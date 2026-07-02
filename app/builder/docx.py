from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.models.translation_unit import TranslationUnit, UnitType


def _get_style_name(document: DocxDocument, preferred: str | None, fallback: str) -> str:
    for style_name in [preferred, fallback]:
        if style_name and style_name in document.styles:
            return style_name
    return "Normal"


def _resolve_text(unit: TranslationUnit) -> str:
    """Prefer translated text when available, falling back to source text."""
    return unit.translated_text if unit.is_translated else unit.text


def _apply_paragraph_formatting(paragraph: object, formatting: dict[str, object]) -> None:
    if not formatting:
        return

    if "bold" in formatting:
        for run in paragraph.runs:
            run.bold = bool(formatting["bold"])
    if "italic" in formatting:
        for run in paragraph.runs:
            run.italic = bool(formatting["italic"])
    if "underline" in formatting:
        for run in paragraph.runs:
            run.underline = bool(formatting["underline"])
    if "font_name" in formatting:
        for run in paragraph.runs:
            run.font.name = str(formatting["font_name"])
    if "font_size" in formatting:
        for run in paragraph.runs:
            run.font.size = formatting["font_size"]
    if "alignment" in formatting:
        alignment_name = str(formatting["alignment"]).upper()
        alignment_map = {
            "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
            "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
            "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
            "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        paragraph.paragraph_format.alignment = alignment_map.get(alignment_name, WD_ALIGN_PARAGRAPH.LEFT)


def _apply_cell_formatting(cell: object, formatting: dict[str, object]) -> None:
    if not formatting:
        return

    paragraph = cell.paragraphs[0] if cell.paragraphs else None
    if paragraph is not None:
        _apply_paragraph_formatting(paragraph, formatting)

    if "vertical_alignment" in formatting:
        alignment_name = str(formatting["vertical_alignment"]).upper()
        alignment_map = {
            "TOP": WD_CELL_VERTICAL_ALIGNMENT.TOP,
            "CENTER": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            "BOTTOM": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
        }
        cell.vertical_alignment = alignment_map.get(alignment_name, WD_CELL_VERTICAL_ALIGNMENT.TOP)

    if "shading" in formatting:
        fill = formatting["shading"].get("fill") if isinstance(formatting["shading"], dict) else None
        if fill:
            tc_pr = cell._tc.get_or_add_tcPr()
            shading = tc_pr.find(qn("w:shd"))
            if shading is None:
                shading = OxmlElement("w:shd")
                tc_pr.append(shading)
            shading.set(qn("w:fill"), str(fill))
            shading.set(qn("w:val"), "clear")

    if "borders" in formatting:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tc_pr.append(borders)
        for border_name, border_values in formatting["borders"].items():
            border = borders.find(qn(f"w:{border_name}"))
            if border is None:
                border = OxmlElement(f"w:{border_name}")
                borders.append(border)
            if "val" in border_values:
                border.set(qn("w:val"), str(border_values["val"]))
            if "size" in border_values:
                border.set(qn("w:sz"), str(border_values["size"]))
            if "color" in border_values:
                border.set(qn("w:color"), str(border_values["color"]))


def _append_unit_to_document(document: DocxDocument, unit: TranslationUnit) -> None:
    if unit.type == UnitType.TABLE_CELL:
        return

    text = _resolve_text(unit)

    if unit.type == UnitType.HEADING:
        level = unit.formatting.get("level")
        if level is None:
            level = unit.metadata.get("level")
        if level is None:
            level = 1
        style_name = _get_style_name(document, f"Heading {level}", "Heading 1")
        document.add_paragraph(text, style=style_name)
        return

    if unit.type == UnitType.LIST_ITEM:
        preferred_style = unit.formatting.get("style_name") or unit.metadata.get("style_name")
        style_name = _get_style_name(document, preferred_style, "List Bullet")
        document.add_paragraph(text, style=style_name)
        return

    if unit.type == UnitType.QUOTE:
        style_name = _get_style_name(document, "Intense Quote", "Quote")
        document.add_paragraph(text, style=style_name)
        return

    if unit.type == UnitType.CODE_BLOCK:
        style_name = _get_style_name(document, "Code", "Normal")
        document.add_paragraph(text, style=style_name)
        return

    if unit.type == UnitType.CAPTION:
        style_name = _get_style_name(document, "Caption", "Normal")
        document.add_paragraph(text, style=style_name)
        return

    document.add_paragraph(text)


def _append_tables_to_document(document: DocxDocument, table_units: list[TranslationUnit]) -> None:
    """Group table-cell units by source table and rebuild each table, restoring merges."""
    tables_by_index: dict[int, list[TranslationUnit]] = {}
    for unit in table_units:
        table_index = unit.location.get("table_index", 0)
        tables_by_index.setdefault(table_index, []).append(unit)

    for table_index in sorted(tables_by_index):
        group = tables_by_index[table_index]

        total_rows = max(
            (unit.location.get("table_rows", 0) for unit in group), default=0
        ) or max(
            unit.location.get("row", 0) + unit.location.get("row_span", 1)
            for unit in group
        )
        total_cols = max(
            (unit.location.get("table_cols", 0) for unit in group), default=0
        ) or max(
            unit.location.get("column", 0) + unit.location.get("column_span", 1)
            for unit in group
        )

        table = document.add_table(rows=total_rows, cols=total_cols)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        for unit in group:
            row = unit.location.get("row", 0)
            column = unit.location.get("column", 0)
            row_span = unit.location.get("row_span", 1)
            column_span = unit.location.get("column_span", 1)

            target_cell = table.cell(row, column)
            if row_span > 1 or column_span > 1:
                bottom_right = table.cell(row + row_span - 1, column + column_span - 1)
                target_cell = target_cell.merge(bottom_right)

            target_cell.text = _resolve_text(unit)
            _apply_cell_formatting(target_cell, unit.formatting)


def build_docx_from_units(units: list[TranslationUnit], output_path: str | Path | None = None) -> Path | bytes:
    """Rebuild a DOCX document from translation units and optionally save it to disk."""
    document = Document()

    paragraph_units = [unit for unit in units if unit.type != UnitType.TABLE_CELL]
    table_units = [unit for unit in units if unit.type == UnitType.TABLE_CELL]

    # Preserve original document order; units missing a paragraph_index sort last.
    paragraph_units.sort(
        key=lambda unit: unit.location.get("paragraph_index", float("inf"))
    )

    for unit in paragraph_units:
        _append_unit_to_document(document, unit)

    if table_units:
        _append_tables_to_document(document, table_units)

    if output_path is None:
        document_bytes = io.BytesIO()
        document.save(document_bytes)
        return document_bytes.getvalue()

    output = Path(output_path)
    document.save(output)
    return output