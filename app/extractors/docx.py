from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

from app.models.translation_unit import TranslationUnit, UnitType


def _classify_paragraph_type(paragraph: object, paragraph_index: int) -> tuple[UnitType, str]:
    style_name = (getattr(paragraph, "style", None).name or "").lower()
    if style_name.startswith("heading"):
        return UnitType.HEADING, f"h{paragraph_index}"
    if "list" in style_name:
        return UnitType.LIST_ITEM, f"li{paragraph_index}"
    if "quote" in style_name:
        return UnitType.QUOTE, f"q{paragraph_index}"
    if "code" in style_name:
        return UnitType.CODE_BLOCK, f"code{paragraph_index}"
    if "caption" in style_name:
        return UnitType.CAPTION, f"cap{paragraph_index}"
    return UnitType.PARAGRAPH, f"p{paragraph_index}"


def _extract_paragraph_formatting(paragraph: object) -> dict[str, object]:
    formatting: dict[str, object] = {}
    runs = list(getattr(paragraph, "runs", []) or [])
    if runs:
        formatting["bold"] = any(getattr(run, "bold", False) for run in runs)
        formatting["italic"] = any(getattr(run, "italic", False) for run in runs)
        formatting["underline"] = any(getattr(run, "underline", False) for run in runs)

        for run in runs:
            font_name = getattr(getattr(run, "font", None), "name", None)
            if font_name:
                formatting["font_name"] = font_name
                break

        for run in runs:
            font_size = getattr(getattr(run, "font", None), "size", None)
            if font_size:
                formatting["font_size"] = font_size
                break

    alignment = getattr(getattr(paragraph, "paragraph_format", None), "alignment", None)
    if alignment is not None:
        formatting["alignment"] = alignment.name

    return formatting


def _extract_cell_formatting(cell: object) -> dict[str, object]:
    formatting = _extract_paragraph_formatting(cell.paragraphs[0]) if cell.paragraphs else {}
    vertical_alignment = getattr(cell, "vertical_alignment", None)
    if vertical_alignment is not None:
        formatting["vertical_alignment"] = vertical_alignment.name

    tc_pr = getattr(cell._tc, "get_or_add_tcPr", lambda: None)()
    if tc_pr is not None:
        shading = tc_pr.find(qn("w:shd"))
        if shading is not None:
            fill = shading.get(qn("w:fill"))
            if fill:
                formatting["shading"] = {"fill": fill}

        borders: dict[str, dict[str, str]] = {}
        for border_name in ("top", "bottom", "left", "right"):
            border = tc_pr.find(qn(f"w:{border_name}"))
            if border is None:
                continue
            border_values: dict[str, str] = {}
            value = border.get(qn("w:val"))
            if value:
                border_values["val"] = value
            size = border.get(qn("w:sz"))
            if size:
                border_values["size"] = size
            color = border.get(qn("w:color"))
            if color:
                border_values["color"] = color
            if border_values:
                borders[border_name] = border_values
        if borders:
            formatting["borders"] = borders

    return formatting


def extract_translation_units_from_docx(file_path: str, language: str = "en") -> list[TranslationUnit]:
    """Extract translatable units from a DOCX file as a list of typed units."""
    document = Document(file_path)
    units: list[TranslationUnit] = []

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        unit_type, unit_id = _classify_paragraph_type(paragraph, paragraph_index)
        units.append(
            TranslationUnit(
                id=unit_id,
                text=text,
                type=unit_type,
                location={"paragraph_index": paragraph_index},
                document_type="docx",
                formatting={
                    **_extract_paragraph_formatting(paragraph),
                    "preserve_whitespace": True,
                    "style_name": getattr(paragraph.style, "name", None),
                },
                metadata={"language": language, "char_count": len(text)},
            )
        )

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                text = " ".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                ).strip()
                if not text:
                    continue

                units.append(
                    TranslationUnit(
                        id=f"t{table_index}_r{row_index}_c{column_index}",
                        text=text,
                        type=UnitType.TABLE_CELL,
                        location={
                            "table_index": table_index,
                            "row": row_index,
                            "column": column_index,
                        },
                        document_type="docx",
                        formatting={
                            **_extract_cell_formatting(cell),
                            "preserve_whitespace": True,
                        },
                        metadata={"language": language, "char_count": len(text)},
                    )
                )

    return units


def extract_text_from_docx(file_path: str, language: str = "en") -> list[str]:
    """Backward-compatible helper returning paragraph text from a DOCX file."""
    return [
        unit.text
        for unit in extract_translation_units_from_docx(file_path, language=language)
        if unit.type == UnitType.PARAGRAPH
    ]
